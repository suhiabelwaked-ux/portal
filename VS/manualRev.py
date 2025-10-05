import fitz  # PyMuPDF
import re
import os
import sys
from docx import Document
from docx.shared import Inches

def find_section_bbox(page: fitz.Page, header_rect: fitz.Rect, start_keyword: str, end_keywords: list):
    """Finds the bounding box of a section on a page for image cropping."""
    try:
        words = page.get_text("words") 
        
        start_word = next((w for w in words if start_keyword in w[4] and w[1] > header_rect.y1), None)
        if not start_word:
            return None

        end_word = None
        min_end_y = float('inf')
        
        for keyword in end_keywords:
            candidates = [w for w in words if keyword in w[4] and w[1] > start_word[1]]
            if candidates:
                first_candidate = min(candidates, key=lambda w: w[1])
                if first_candidate[1] < min_end_y:
                    min_end_y = first_candidate[1]
                    end_word = first_candidate
        
        page_rect = page.rect
        x0 = page_rect.x0 + 30
        top = start_word[1] - 10
        x1 = page_rect.x1 - 30
        bottom = end_word[1] - 10 if end_word else page_rect.y1 - 30
        
        if bottom <= top:
            bottom = page_rect.y1 - 30

        if end_word and (end_word[1] - start_word[1] < 20):
            return None

        return fitz.Rect(x0, top, x1, bottom)
    except Exception:
        return None

def process_cis_benchmark_to_word(pdf_path: str, output_docx_path: str):
    """Parses a CIS Benchmark PDF and generates a summary Word document."""
    print(f"--- Starting processing for '{pdf_path}' ---")
    try:
        doc = fitz.Document(pdf_path)
        full_text = "".join(page.get_text() for page in doc)
        print("✅ PDF opened and text extracted successfully.")
    except Exception as e:
        print(f"❌ ERROR: Failed to open or read PDF file: {e}")
        return
    
    image_dir = os.path.join(os.path.dirname(output_docx_path), "audit_screenshots")
    if not os.path.exists(image_dir): os.makedirs(image_dir)
        
    word_doc = Document()
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    word_doc.add_heading(f"CIS Benchmark Summary: {base_name}", level=0)

    header_regex = re.compile(r"^\s*(\d+(?:\.\d+)+\s+.*?)\s*$", re.MULTILINE)
    headers = list(header_regex.finditer(full_text))
    print(f"✅ Found {len(headers)} recommendations. Processing...")

    for i, current_match in enumerate(headers):
        header_text = current_match.group(1).strip()
        print(f"\n--- Processing Recommendation: {header_text} ---")

        content_start_pos = current_match.end()
        content_end_pos = headers[i + 1].start() if i + 1 < len(headers) else len(full_text)
        
        section_content = full_text[content_start_pos:content_end_pos]
            
        desc_match = re.search(r"Description:\s*(.*?)(?=Rationale:|Audit:|Remediation:|$)", section_content, re.DOTALL)
        remed_match = re.search(r"Remediation:\s*(.*?)(?=Default Value:|References:|$)", section_content, re.DOTALL)
        has_audit = "Audit:" in section_content

        if desc_match or remed_match or has_audit:
            word_doc.add_heading(header_text, level=2)

            if desc_match and desc_match.group(1).strip():
                word_doc.add_heading("Description", level=3)
                word_doc.add_paragraph(desc_match.group(1).strip())

            if has_audit:
                word_doc.add_heading("Audit", level=3)
                image_generated = False
                for page in doc:
                    header_rects = page.search_for(header_text.splitlines()[0])
                    if header_rects:
                        end_keywords = ["Remediation:", "Default Value:", "References:", "CIS Controls:", "Additional Information:"]
                        bbox = find_section_bbox(page, header_rects[0], "Audit:", end_keywords)
                        
                        if bbox:
                            filename_safe_header = re.sub(r'[^\w.-]', '_', header_text.split(' ')[0])
                            img_file = os.path.join(image_dir, f"audit_{filename_safe_header}.png")
                            
                            pix = page.get_pixmap(clip=bbox, dpi=200)
                            pix.save(img_file)
                            
                            try:
                                word_doc.add_picture(img_file, width=Inches(6.0))
                                image_generated = True
                            except Exception as e:
                                word_doc.add_paragraph(f"[ERROR embedding image: {e}]")
                            break
                
                if not image_generated:
                    word_doc.add_paragraph("[Could not generate screenshot. The section might be empty or formatted unusually.]")
            
            if remed_match and remed_match.group(1).strip():
                word_doc.add_heading("Remediation", level=3)
                word_doc.add_paragraph(remed_match.group(1).strip())

    doc.close()
    word_doc.save(output_docx_path)
    print("\n" + "="*80)
    print(f"✅ Success! Summary saved to: {output_docx_path}")
    print(f"📷 Screenshots saved in: {image_dir}")
    print("="*80)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python your_script_name.py <path_to_your_pdf_file.pdf>")
        sys.exit(1)
    
    pdf_file_path = sys.argv[1]
    if not os.path.exists(pdf_file_path):
        print(f"❌ Error: The file '{pdf_file_path}' was not found.")
        sys.exit(1)
        
    try:
        process_cis_benchmark_to_word(pdf_file_path, 'output.docx')
    except Exception as e:
        print(f"\n❌ An unexpected error occurred during processing: {e}")