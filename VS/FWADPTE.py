import re
import pypdf
import argparse
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import mysql.connector
from mysql.connector import Error
import shutil

# --- CONFIGURATION ---
CONFIG = {
    'fonts': {
        'body': 'Aptos',
        'heading': 'Aptos SemiBold',
    },
    'colors': {
        'text': RGBColor(0, 0, 0),
        'heading': RGBColor(37, 55, 65),
        'severity': {
            'Critical': RGBColor(192, 0, 0),
            'High': RGBColor(255, 0, 0),
            'Medium': RGBColor(255, 192, 0),
            'Low': RGBColor(146, 205, 80),
        }
    },
    'company_name': 'FutureTEC',
    'source_name': 'Nipper'
}

# Pre-compiled regular expressions for efficiency
PATTERNS = {
    'start_marker': "2 Security Audit",
    'header': re.compile(r"^\s*(2\.\d+)\s+.*$"),
    'finding_subheader': re.compile(r"^\s*\d\.\d+\.\d+\s+Finding"),
    'reco_subheader': re.compile(r"^\s*\d\.\d+\.\d+\s+Recommendation"),
    'affected_subheader': re.compile(r"^\s*\d\.\d+\.\d+\s+Affected Device"),
    'stop_subheader': re.compile(r"^\s*\d\.\d+\.\d+\s+(Impact|Ease)"),
    'sidebar_metadata': re.compile(r"\s*(Overall:|Impact:|Ease:|Fix:|Type:|Finding ID:).*", re.IGNORECASE),
    'junk_line': re.compile(
        r"Table\s+\d+|detailed in Table|Go to the report contents|Notes\s+for|Severity|Host\s+IP|Port\s*/\s*Protocol|Vulnerability|^\s*Line\s+Access\s+Login|Interface\s+Active\s+Description"
    )
}

# --- Database Configuration (REMOVED HARDCODED CREDENTIALS) ---
# DB_CONFIG = {  # UNUSED - Function uses SQLite below
    # 'user': 'root',  # REMOVED HARDCODED USER
    # 'password': 'admin_2025',  # REMOVED HARDCODED PASSWORD
    # 'host': 'localhost',  # REMOVED HARDCODED HOST
    # 'database': 'vulnerability_db'  # REMOVED HARDCODED DATABASE
# } # END UNUSED DB_CONFIG

# --- Database Connection and Severity Lookup ---
def get_severity_from_db(finding_title):
    try:
        # Get MySQL connection parameters from environment or use defaults
        import os
        config = {
            'host': os.getenv('MYSQL_HOST', 'localhost'),
            'database': os.getenv('MYSQL_DATABASE', 'vulnerability_db'),
            'user': os.getenv('MYSQL_USER', 'root'),
            'password': os.getenv('MYSQL_PASSWORD', ''),
            'port': int(os.getenv('MYSQL_PORT', 3306))
        }
        
        conn = mysql.connector.connect(**config)
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT severity FROM firewall_vulnerabilities WHERE name = %s", (finding_title,))
        result = cursor.fetchone()
        
        if not result:
            print(f"Vulnerability '{finding_title}' not found. Adding to pending table.")
            cursor.execute("INSERT INTO pending (name, severity) VALUES (%s, %s) ON DUPLICATE KEY UPDATE severity=%s", (finding_title, "Pending", "Pending"))
            conn.commit()
            return None
        
        return result.get('severity') if result else None
    
    except Error as err:
        print(f"Database error: {err}")
        return None
    finally:
        if 'conn' in locals() and conn.is_connected():
            cursor.close()
            conn.close()

class ReportGenerator:
    """
    A class to extract security findings from a PDF and generate a styled Word document.
    """
    def __init__(self, pdf_path, output_docx):
        self.pdf_path = pdf_path
        self.output_docx = output_docx
        self.doc = Document()
        self.skipped_sections = []

    def run(self):
        """Executes the entire process of PDF parsing and DOCX generation."""
        try:
            print(f"Reading PDF: {self.pdf_path}...")
            full_text = self._extract_text_from_pdf()
            
            print("Parsing findings from text...")
            findings = self._parse_findings_from_text(full_text)

            print("Categorizing findings...")
            audit_findings, policy_findings = self._categorize_findings(findings)

            print("Writing findings to Word document...")
            self._write_sections_to_doc(audit_findings, policy_findings)

            self.doc.save(self.output_docx)
            print(f"\n✅ Successfully created final Word document: {self.output_docx}")
            return True

        except FileNotFoundError:
            print(f"❌ Error: The file '{self.pdf_path}' was not found.")
            return False
        except Exception as e:
            print(f"❌ An unexpected error occurred: {e}")
            return False

    def _extract_text_from_pdf(self):
        """Opens and extracts all text from the source PDF file."""
        pdf_reader = pypdf.PdfReader(self.pdf_path)
        return "".join(page.extract_text() + "\n" for page in pdf_reader.pages)

    def _parse_findings_from_text(self, full_text):
        """Parses the raw text to identify and structure individual findings."""
        lines = full_text.splitlines()
        findings = []
        current_finding_data = {}
        in_audit_section = False
        current_state = None  # Can be 'finding', 'reco', or 'affected'

        for line in lines:
            if not in_audit_section:
                if PATTERNS['start_marker'] in line:
                    in_audit_section = True
                continue

            if PATTERNS['header'].match(line):
                if current_finding_data:
                    findings.append(current_finding_data)
                current_finding_data = {'header': self._clean_text(line), 'finding_text': '', 'reco_text': '', 'affected_text': ''}
                current_state = None
                continue

            if PATTERNS['junk_line'].search(line) or PATTERNS['stop_subheader'].match(line):
                current_state = None
                continue
            elif PATTERNS['affected_subheader'].match(line):
                current_state = 'affected'
                continue
            elif PATTERNS['reco_subheader'].match(line):
                current_state = 'reco'
                continue
            elif PATTERNS['finding_subheader'].match(line):
                current_state = 'finding'
                continue

            cleaned_line = PATTERNS['sidebar_metadata'].sub("", line).strip()
            if cleaned_line and current_state:
                current_finding_data[f'{current_state}_text'] += cleaned_line + " "
        
        if current_finding_data:
            findings.append(current_finding_data)
            
        return findings

    def _categorize_findings(self, findings):
        """Sorts findings into 'Audit' and 'Policy' categories."""
        audit_findings, policy_findings = [], []
        skip_titles = {
            "rules allow access to potentially unnecessary services",
            "rules allow access to potentially sensitive services",
            "No Network Filtering Rules Were Configured"
        }
        for finding in findings:
            title_only = re.sub(r"^\s*\d+(\.\d+)*\s*", "", finding['header']).strip().lower()
            if title_only in skip_titles:
                continue
            elif title_only.startswith(('rules', 'filter')):
                policy_findings.append(finding)
            else:
                audit_findings.append(finding)
        return audit_findings, policy_findings

    def _write_sections_to_doc(self, audit_findings, policy_findings):
        self.skipped_sections = []
        self._style_main_heading('2.Security Audit')
        count = 1
        for finding in audit_findings:
            if self._write_finding_to_doc(finding, count):
                count += 1
        
        if policy_findings:
            self.doc.add_page_break()
            self._style_main_heading('3.Policy Audit')
            count = 1
            for finding in policy_findings:
                if self._write_finding_to_doc(finding, count):
                    count += 1
    
    def _write_finding_to_doc(self, finding, count):
        header = finding['header']
        finding_text = self._clean_text(finding['finding_text'])
        reco_text = self._clean_text(finding['reco_text'])
        affected_text = self._clean_text(finding.get('affected_text', ''))

        if not (header and finding_text and reco_text):
            if header and header not in self.skipped_sections:
                self.skipped_sections.append(header)
            return False

        if count > 1:
            self.doc.add_paragraph("")

        title_only = re.sub(r"^\s*\d+(\.\d+)*\s*", "", header).strip()
        title_para = self.doc.add_paragraph()
        title_para.add_run(f"{count}-  {title_only}")
        self._style_paragraph(
            title_para, font_name=CONFIG['fonts']['heading'], size=12, bold=True,
            font_color=CONFIG['colors']['heading'], style='Heading 2'
        )
       
        self._write_severity_block(title_only, count)
        sub_section_count = 2

        if affected_text:
            self._write_affected_block(count, sub_section_count, affected_text)
            sub_section_count += 1

        self._write_description_block(count, sub_section_count, finding_text)
        sub_section_count += 1

        self._write_recommendation_block(count, sub_section_count, reco_text)

        return True

    def _write_severity_block(self, title_only, count):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        
        font_name = CONFIG['fonts']['heading']
        font_size = Pt(11)
        heading_color = CONFIG['colors']['heading']

        run_intro = p.add_run(f"{count}.1 Severity: ")
        run_intro.bold = True
        run_intro.font.name = font_name
        run_intro.font.size = font_size
        run_intro.font.color.rgb = heading_color

        specific_severity = get_severity_from_db(title_only)
        severity_colors = CONFIG['colors']['severity']

        if specific_severity and specific_severity in severity_colors:
            run_level = p.add_run(specific_severity)
            run_level.font.color.rgb = severity_colors[specific_severity]
            run_level.bold = True
            run_level.font.name = font_name
            run_level.font.size = font_size
        else:
            all_severities = [
                ('Critical', RGBColor(192, 0, 0)),
                ('High', RGBColor(255, 0, 0)),
                ('Medium', RGBColor(255, 192, 0)),
                ('Low', RGBColor(146, 205, 80))
            ]
            for i, (label, color) in enumerate(all_severities):
                if i > 0:
                    sep = p.add_run(' / ')
                    sep.bold = True
                    sep.font.name = font_name
                    sep.font.size = font_size
                    sep.font.color.rgb = heading_color
                r = p.add_run(label)
                r.font.name = font_name
                r.font.size = font_size
                r.bold = True
                r.font.color.rgb = color
        
        p.paragraph_format.space_after = Pt(0)

    def _write_affected_block(self, count, sub_count, text):
        self._write_subsection_heading(f"{count}.{sub_count} Affected Device:")
        parts = text.split(';')
        for part in parts:
            if part.strip():
                self._write_bullet_point(part.strip())

    def _write_description_block(self, count, sub_count, text):
        text = text.replace(CONFIG['source_name'], CONFIG['company_name'])
        
        truncation_phrases = [
            r"These are detailed below", r"Those filter rule lists are detailed below",
            r"Those filter rules are listed below", r"this filter rules are listed below",
            r"Those filter rules are listedbelow"
        ]
        pattern = re.compile("|".join(truncation_phrases), re.IGNORECASE)
        text_to_write = pattern.split(text, maxsplit=1)[0].strip()
        
        self._write_subsection_heading(f"{count}.{sub_count} Description:")

        company_name = re.escape(CONFIG['company_name'])
        
        identified_pattern = re.compile(f"{company_name} identified", re.IGNORECASE)
        determined_pattern = re.compile(f"{company_name} determined", re.IGNORECASE)
        
        identified_match = identified_pattern.search(text_to_write)
        determined_match = determined_pattern.search(text_to_write)

        if identified_match:
            actual_trigger_phrase = identified_match.group(0)
            text_parts = identified_pattern.split(text_to_write, maxsplit=1)
            
            if text_parts[0].strip():
                self._write_bullet_point(text_parts[0].strip())
            
            self._write_bullet_point(actual_trigger_phrase)
            
            if len(text_parts) > 1 and text_parts[1].strip():
                self._write_bullet_point(text_parts[1].strip())

        elif determined_match:
            actual_trigger_phrase = determined_match.group(0)
            text_parts = determined_pattern.split(text_to_write, maxsplit=1)

            if text_parts[0].strip():
                self._write_bullet_point(text_parts[0].strip())
            
            text_after = text_parts[1].strip() if len(text_parts) > 1 else ""
            combined_text = f"{actual_trigger_phrase} {text_after}".strip()
            self._write_bullet_point(combined_text)
            
        else:
            self._write_bullet_point(text_to_write)

    def _write_recommendation_block(self, count, sub_count, text):
        self._write_subsection_heading(f"{count}.{sub_count} Recommendation:")
        
        text = text.replace(CONFIG['source_name'], CONFIG['company_name'])
        trigger_phrase = f"{CONFIG['company_name']} recommends that:"

        if trigger_phrase in text:
            parts = text.split(trigger_phrase, 1)
            if parts[0].strip():
                self._write_bullet_point(parts[0].strip())
            
            self._write_bullet_point(trigger_phrase)

            recommendations = parts[1].split(';')
            for item in recommendations:
                if item.strip():
                    final_text = f"◦   {item.strip()}"
                    self._write_bullet_point(final_text, use_native_bullet=False)
        else:
            self._write_bullet_point(text)

    def _write_subsection_heading(self, text):
        p = self.doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        self._style_paragraph(p, font_name=CONFIG['fonts']['heading'], size=11, bold=True, font_color=CONFIG['colors']['heading'])

    def _write_bullet_point(self, text, use_native_bullet=True):
        style = 'List Bullet' if use_native_bullet else None
        p = self.doc.add_paragraph(text, style=style)
        p.paragraph_format.left_indent = Inches(0.75)
        self._style_paragraph(p)
    
    def _style_main_heading(self, text):
        heading = self.doc.add_heading(text, level=1)
        heading.runs[0].font.size = Pt(14.5)
        heading.runs[0].font.name = CONFIG['fonts']['heading']
        heading.runs[0].bold = True
        heading.runs[0].font.color.rgb = CONFIG['colors']['heading']
        heading.paragraph_format.space_after = Pt(6)

        pPr = heading._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '242C3F')
        
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _style_paragraph(self, paragraph, font_name=None, size=11, bold=False, spacing_after=0, font_color=None, style=None):
        if font_name is None: font_name = CONFIG['fonts']['body']
        if font_color is None: font_color = CONFIG['colors']['text']
        
        if style:
            paragraph.style = style
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = font_color

        paragraph.paragraph_format.space_after = Pt(spacing_after)

    def _clean_text(self, text):
        if not text: return ""
        return ' '.join(text.strip().split())


def extract_and_style_audit_findings(pdf_path, output_docx):
    """
    A class to extract security findings from a PDF and generate a styled Word document.
    """
    try:
        print(f"Reading PDF: {pdf_path}...")
        full_text = _extract_text_from_pdf(pdf_path)
        
        print("Parsing findings from text...")
        findings = _parse_findings_from_text(full_text)

        print("Categorizing findings...")
        audit_findings, policy_findings = _categorize_findings(findings)

        print("Writing findings to Word document...")
        doc = Document()
        _write_sections_to_doc(doc, audit_findings, policy_findings)
        
        doc.save(output_docx)
        print(f"\n✅ Successfully created final Word document: {output_docx}")
        return True

    except FileNotFoundError:
        print(f"❌ Error: The file '{pdf_path}' was not found.")
        return False
    except Exception as e:
        print(f"❌ An unexpected error occurred: {e}")
        return False

def _extract_text_from_pdf(pdf_path):
    """Opens and extracts all text from the source PDF file."""
    pdf_reader = pypdf.PdfReader(pdf_path)
    return "".join(page.extract_text() + "\n" for page in pdf_reader.pages)

def _parse_findings_from_text(full_text):
    """Parses the raw text to identify and structure individual findings."""
    lines = full_text.splitlines()
    findings = []
    current_finding_data = {}
    in_audit_section = False
    current_state = None  # Can be 'finding', 'reco', or 'affected'

    for line in lines:
        if not in_audit_section:
            if PATTERNS['start_marker'] in line:
                in_audit_section = True
            continue

        if PATTERNS['header'].match(line):
            if current_finding_data:
                findings.append(current_finding_data)
            current_finding_data = {'header': clean_text(line), 'finding_text': '', 'reco_text': '', 'affected_text': ''}
            current_state = None
            continue

        if PATTERNS['junk_line'].search(line) or PATTERNS['stop_subheader'].match(line):
            current_state = None
            continue
        elif PATTERNS['affected_subheader'].match(line):
            current_state = 'affected'
            continue
        elif PATTERNS['reco_subheader'].match(line):
            current_state = 'reco'
            continue
        elif PATTERNS['finding_subheader'].match(line):
            current_state = 'finding'
            continue

        cleaned_line = PATTERNS['sidebar_metadata'].sub("", line).strip()
        if cleaned_line and current_state:
            current_finding_data[f'{current_state}_text'] += cleaned_line + " "
    
    if current_finding_data:
        findings.append(current_finding_data)
        
    return findings

def _categorize_findings(findings):
    """Sorts findings into 'Audit' and 'Policy' categories."""
    audit_findings, policy_findings = [], []
    skip_titles = {
        "rules allow access to potentially unnecessary services",
        "rules allow access to potentially sensitive services",
        "No Network Filtering Rules Were Configured"
    }
    for finding in findings:
        title_only = re.sub(r"^\s*\d+(\.\d+)*\s*", "", finding['header']).strip().lower()
        if title_only in skip_titles:
            continue
        elif title_only.startswith(('rules', 'filter')):
            policy_findings.append(finding)
        else:
            audit_findings.append(finding)
    return audit_findings, policy_findings

def _write_sections_to_doc(doc, audit_findings, policy_findings):
    """Writes all the categorized findings into the document."""
    skipped_sections = []
    
    _style_main_heading(doc, '2.Security Audit')
    count = 1
    for finding in audit_findings:
        if _write_finding_to_doc(doc, finding, count):
            count += 1
    
    if policy_findings:
        doc.add_page_break()
        _style_main_heading(doc, '3.Policy Audit')
        count = 1
        for finding in policy_findings:
            if _write_finding_to_doc(doc, finding, count):
                count += 1
    
def _style_main_heading(doc, text):
    heading = doc.add_heading(text, level=1)
    run = heading.runs[0]
    font = run.font
    font.size = Pt(14.5)
    font.name = 'Aptos SemiBold'
    font.bold = True
    font.color.rgb = RGBColor(37, 55, 65)
    heading.paragraph_format.space_after = Pt(6)

def _write_affected_block(doc, count, sub_count, text):
    _write_subsection_heading(doc, f"{count}.{sub_count} Affected Device:")
    parts = text.split(';')
    for part in parts:
        if part.strip():
            _write_bullet_point(doc, part.strip())

def _write_description_block(doc, count, sub_count, text):
    text = text.replace(CONFIG['source_name'], CONFIG['company_name'])
    
    truncation_phrases = [
        r"These are detailed below", r"Those filter rule lists are detailed below",
        r"Those filter rules are listed below", r"this filter rules are listed below",
        r"Those filter rules are listedbelow"
    ]
    pattern = re.compile("|".join(truncation_phrases), re.IGNORECASE)
    text_to_write = pattern.split(text, maxsplit=1)[0].strip()
    
    _write_subsection_heading(doc, f"{count}.{sub_count} Description:")

    company_name = re.escape(CONFIG['company_name'])
    
    identified_pattern = re.compile(f"{company_name} identified", re.IGNORECASE)
    determined_pattern = re.compile(f"{company_name} determined", re.IGNORECASE)
    
    identified_match = identified_pattern.search(text_to_write)
    determined_match = determined_pattern.search(text_to_write)

    if identified_match:
        actual_trigger_phrase = identified_match.group(0)
        text_parts = identified_pattern.split(text_to_write, maxsplit=1)
        
        if text_parts[0].strip():
            _write_bullet_point(doc, text_parts[0].strip())
        
        _write_bullet_point(doc, actual_trigger_phrase)
        
        if len(text_parts) > 1 and text_parts[1].strip():
            _write_bullet_point(doc, text_parts[1].strip())

    elif determined_match:
        actual_trigger_phrase = determined_match.group(0)
        text_parts = determined_pattern.split(text_to_write, maxsplit=1)

        if text_parts[0].strip():
            _write_bullet_point(doc, text_parts[0].strip())
        
        text_after = text_parts[1].strip() if len(text_parts) > 1 else ""
        combined_text = f"{actual_trigger_phrase} {text_after}".strip()
        _write_bullet_point(doc, combined_text)
        
    else:
        _write_bullet_point(doc, text_to_write)

def _write_recommendation_block(doc, count, sub_count, text):
    _write_subsection_heading(doc, f"{count}.{sub_count} Recommendation:")
    
    text = text.replace(CONFIG['source_name'], CONFIG['company_name'])
    trigger_phrase = f"{CONFIG['company_name']} recommends that:"

    if trigger_phrase in text:
        parts = text.split(trigger_phrase, 1)
        if parts[0].strip():
            _write_bullet_point(doc, parts[0].strip())
        
        _write_bullet_point(doc, trigger_phrase)

        recommendations = parts[1].split(';')
        for item in recommendations:
            if item.strip():
                final_text = f"◦   {item.strip()}"
                _write_bullet_point(doc, final_text, use_native_bullet=False)
        else:
            _write_bullet_point(doc, text)

def _write_subsection_heading(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(0)
    _style_paragraph(doc, p, font_name=CONFIG['fonts']['heading'], size=11, bold=True, font_color=CONFIG['colors']['heading'])

def _write_bullet_point(doc, text, use_native_bullet=True):
    style = 'List Bullet' if use_native_bullet else None
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.left_indent = Inches(0.75)
    _style_paragraph(doc, p)

def _style_paragraph(doc, paragraph, font_name=None, size=11, bold=False, spacing_after=0, font_color=None, style=None):
    if font_name is None: font_name = CONFIG['fonts']['body']
    if font_color is None: font_color = CONFIG['colors']['text']
    
    if style:
        paragraph.style = style
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = font_color

    paragraph.paragraph_format.space_after = Pt(spacing_after)

def _clean_text(self, text):
    if not text: return ""
    return ' '.join(text.strip().split())


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Extracts and styles security audit findings from a PDF into a Word document."
    )
    parser.add_argument("pdf_path", help="The full path to the input PDF file.")
    parser.add_argument(
        "-o",
        "--output",
        help="The path for the output Word document. Defaults to the same name as the input PDF with a .docx extension.",
    )
    args = parser.parse_args()

    if args.output:
        output_path = args.output
    else:
        pdf_directory = os.path.dirname(args.pdf_path)
        pdf_basename = os.path.splitext(os.path.basename(args.pdf_path))[0]
        output_path = os.path.join(pdf_directory, f"{pdf_basename}.docx")

    generator = ReportGenerator(args.pdf_path, output_path)
    generator.run()