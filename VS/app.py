import os
import time
# MySQL database integration
import pymysql
from flask import Flask, render_template, request, redirect, url_for, session, send_file, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
from flask_wtf.csrf import CSRFProtect
from typing import List, Dict, Any
import shutil
import uuid
import json
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
try:
    import fitz # PyMuPDF
except ImportError:
    print("Warning: PyMuPDF not available. PDF processing may be limited.")
    fitz = None
import re
import base64
import pypdf
import tempfile
import sys
from database import get_db_connection
# Cross-platform file lock/unlock functions
if sys.platform == "win32":
    import msvcrt

    def lock_file(file, mode):
        if mode == "r":  # shared lock
            msvcrt.locking(file.fileno(), msvcrt.LK_RLCK, 1)
        elif mode == "w":  # exclusive lock
            msvcrt.locking(file.fileno(), msvcrt.LK_LOCK, 1)

    def unlock_file(file):
        try:
            msvcrt.locking(file.fileno(), msvcrt.LK_UNLCK, 1)
        except OSError:
            pass

else:
    import fcntl

    def lock_file(file, mode):
        if mode == "r":
            lock_file(file, "r")
        elif mode == "w":
            lock_file(file, "w")

    def unlock_file(file):
        lock_file(file, unlock_file)


# File processing imports
import TestwithSeverity
import FWADPTE

# --- Flask App Initialization ---
app = Flask(__name__)
app.secret_key = os.urandom(24)

# Configure session for iframe compatibility (Replit requirement)
app.config['SESSION_COOKIE_SECURE'] = False
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Enable CORS for iframe compatibility
CORS(app, supports_credentials=True)

# Define upload and output folders
UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
OUTPUT_FOLDER = os.path.join(os.getcwd(), 'output')
TEMP_DATA_FOLDER = os.path.join(os.getcwd(), 'temp_data')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(TEMP_DATA_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TEMP_DATA_FOLDER'] = TEMP_DATA_FOLDER

# Initialize CSRF Protection
csrf = CSRFProtect(app)

# Add CSRF token to template context
@app.context_processor
def inject_csrf_token():
    from flask_wtf.csrf import generate_csrf
    return dict(csrf_token=generate_csrf)

# --- MySQL Database Configuration ---
DB_CONFIG = {
    'user': os.environ.get('DB_USER', 'root'),
    'password': os.environ.get('DB_PASSWORD', 'admin_2025'),
    'host': os.environ.get('DB_HOST', 'localhost'),
    'database': os.environ.get('DB_DATABASE', 'vulnerability_db')
}


def user_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        if not session.get('user_logged_in'):
            # Check if this is an AJAX request expecting JSON response
            if request.is_json or request.headers.get('Content-Type') == 'application/json':
                return jsonify({'ok': False, 'error': 'Session expired. Please re-login.'}), 401
            # For regular requests, redirect to login page
            return redirect(url_for('user_login'))
        return func(*args, **kwargs)
    return wrapper

def admin_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return redirect(url_for('admin_login'))
        return func(*args, **kwargs)
    return wrapper

# --- RBAC (Role-Based Access Control) Functions ---

def get_user_permissions(username):
    """Get all permissions for a user through their group memberships"""
    conn = get_db_connection()
    if not conn:
        return []
    
    cursor = conn.cursor()
    cursor.execute("""
        SELECT DISTINCT p.name as permission_name, p.resource, p.action
        FROM users u
        JOIN user_groups ug ON u.id = ug.user_id
        JOIN `groups` g ON ug.group_id = g.id
        JOIN group_permissions gp ON g.id = gp.group_id
        JOIN permissions p ON gp.permission_id = p.id
        WHERE u.username = %s AND u.status = 'enabled'
    """, (username,))
    
    permissions = cursor.fetchall()
    cursor.close()
    conn.close()
    
    return [perm['permission_name'] for perm in permissions]

def get_user_role(username):
    """Get the primary role of a user"""
    conn = get_db_connection()
    if not conn:
        return None
    
    cursor = conn.cursor()
    cursor.execute("SELECT role FROM users WHERE username = %s", (username,))
    result = cursor.fetchone()
    cursor.close()
    conn.close()
    
    return result['role'] if result else None

def has_permission(username, permission_name):
    """Check if user has a specific permission"""
    return permission_name in get_user_permissions(username)

def require_permission(permission_name):
    """Decorator to require a specific permission"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            # Check if user or admin is logged in
            user_logged_in = session.get('user_logged_in')
            admin_logged_in = session.get('admin_logged_in')
            
            if not user_logged_in and not admin_logged_in:
                if request.is_json or request.headers.get('Content-Type') == 'application/json':
                    return jsonify({'ok': False, 'error': 'Session expired. Please re-login.'}), 401
                return redirect(url_for('user_login'))
            
            # Admin users have all permissions
            if admin_logged_in:
                return func(*args, **kwargs)
            
            # For regular users, check permissions
            username = session.get('username')
            if not username:
                return redirect(url_for('user_login'))
            
            # Check permission
            if not has_permission(username, permission_name):
                if request.is_json or request.headers.get('Content-Type') == 'application/json':
                    return jsonify({'ok': False, 'error': f'Permission denied: {permission_name} required'}), 403
                return render_template('error.html', 
                                     error=f'Access denied: You need {permission_name} permission to access this page.')
            
            return func(*args, **kwargs)
        return wrapper
    return decorator

def require_role(required_role):
    """Decorator to require a specific role"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            # Check if user or admin is logged in
            user_logged_in = session.get('user_logged_in')
            admin_logged_in = session.get('admin_logged_in')
            
            if not user_logged_in and not admin_logged_in:
                if request.is_json or request.headers.get('Content-Type') == 'application/json':
                    return jsonify({'ok': False, 'error': 'Session expired. Please re-login.'}), 401
                return redirect(url_for('user_login'))
            
            # Admin users have all roles
            if admin_logged_in:
                return func(*args, **kwargs)
            
            # For regular users, check role
            username = session.get('username')
            if not username:
                return redirect(url_for('user_login'))
            
            # Check role
            user_role = get_user_role(username)
            if user_role != required_role:
                if request.is_json or request.headers.get('Content-Type') == 'application/json':
                    return jsonify({'ok': False, 'error': f'Role denied: {required_role} role required'}), 403
                return render_template('error.html', 
                                     error=f'Access denied: You need {required_role} role to access this page.')
            
            return func(*args, **kwargs)
        return wrapper
    return decorator

# --- Application Routes ---

@app.route('/')
def home():
    return redirect(url_for('user_login'))

@app.route('/user_login', methods=['GET', 'POST'])
def user_login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        conn = get_db_connection()
        if not conn:
            return "Database connection failed.", 500
        
        cursor = conn.cursor()
        cursor.execute("SELECT id, password_hash, status FROM users WHERE username = %s", (username,))
        user_data = cursor.fetchone()
        cursor.close()
        conn.close()

        if user_data and check_password_hash(user_data['password_hash'], password):
            if user_data['status'] != 'enabled':
                return render_template('user_login.html', error='Account is disabled. Please contact administrator.')
            
            session['user_logged_in'] = True
            session['username'] = username
            session['user_id'] = user_data['id']
            return redirect(url_for('main_page'), code=303)
        else:
            return render_template('user_login.html', error='Invalid credentials')
    return render_template('user_login.html')

@app.route('/main_page', methods=['GET'])
@user_required
def main_page():
    # Get user permissions for conditional UI
    user_permissions = get_user_permissions(session.get('username'))
    return render_template('main_page.html', user_permissions=user_permissions)


@app.route('/lookup_router', methods=['GET', 'POST'])
@user_required
def lookup_router():
    vulnerability_name = None
    severity = None
    not_found = False
    is_pending = False
    
    if request.method == 'POST':
        vulnerability_name = request.form['search_name']
        
        conn = get_db_connection()
        if not conn:
            return redirect(url_for('lookup_router', message="Database connection failed."))
        
        cursor = conn.cursor()
        
        cursor.execute('SELECT severity FROM vulnerabilities WHERE name = %s', (vulnerability_name,))
        query = cursor.fetchone()
        
        if query:
            severity = query['severity']
        else:
            cursor.execute('SELECT severity FROM pending WHERE name = %s', (vulnerability_name,))
            pending_query = cursor.fetchone()
            if pending_query:
                is_pending = True
            else:
                not_found = True
        
        cursor.close()
        conn.close()
    
    return render_template('lookup_router.html',
                           vulnerability_name=vulnerability_name,
                           severity=severity,
                           not_found=not_found,
                           is_pending=is_pending)


@app.route('/lookup_firewall', methods=['GET', 'POST'])
@user_required
def lookup_firewall():
    vulnerability_name = None
    severity = None
    not_found = False
    is_pending = False
    
    if request.method == 'POST':
        vulnerability_name = request.form['search_name']
        
        conn = get_db_connection()
        if not conn:
            return redirect(url_for('lookup_firewall', message="Database connection failed."))
        
        cursor = conn.cursor()
        
        cursor.execute('SELECT severity FROM firewall_vulnerabilities WHERE name = %s', (vulnerability_name,))
        query = cursor.fetchone()

        if query:
            severity = query['severity']
        else:
            cursor.execute('SELECT severity FROM pending WHERE name = %s', (vulnerability_name,))
            pending_query = cursor.fetchone()
            if pending_query:
                is_pending = True
            else:
                not_found = True
        
        cursor.close()
        conn.close()
    
    return render_template('lookup_firewall.html',
                           vulnerability_name=vulnerability_name,
                           severity=severity,
                           not_found=not_found,
                           is_pending=is_pending)

@app.route('/add_to_pending', methods=['POST'])
@user_required
def add_to_pending():
    new_name = request.form['add_name']
    new_severity = request.form['add_severity']
    search_type = request.form.get('search_type')
    submitted_by = session.get('username')

    conn = get_db_connection()
    if not conn:
        return "Database connection failed.", 500
        
    cursor = conn.cursor()
    try:
        cursor.execute('INSERT INTO pending (name, severity, search_type, submitted_by) VALUES (%s, %s, %s, %s)', (new_name, new_severity, search_type, submitted_by))
        conn.commit()
        message = 'Your submission is pending admin approval.'
    except sqlite3.Error as err:
        message = f"An error occurred: {str(err)}"
    finally:
        if 'conn' in locals():
            cursor.close()
            conn.close()
    
    if search_type == 'firewall':
      return redirect(url_for('lookup_firewall', message=message))
    else:
      return redirect(url_for('lookup_router', message=message))

@app.route('/convert_router_page')
@user_required
def convert_router_page():
    message = request.args.get('message')
    download_link = request.args.get('download_link')
    return render_template('convert_router.html', message=message, download_link=download_link)

@app.route('/convert_firewall_page')
@user_required
def convert_firewall_page():
    message = request.args.get('message')
    download_link = request.args.get('download_link')
    return render_template('convert_firewall.html', message=message, download_link=download_link)

@app.route('/convert_router', methods=['POST'])
@user_required
def convert_router():
    message = None
    download_link = None
    
    if 'pdf_file' not in request.files:
        message = 'No file part in the request.'
    else:
        file = request.files['pdf_file']
        if file.filename == '' or not file.filename.endswith('.pdf'):
            message = 'Please select a PDF file.'
        else:
            try:
                filename = secure_filename(file.filename)
                pdf_path = os.path.join(UPLOAD_FOLDER, filename)
                file.save(pdf_path)
                
                output_filename = os.path.splitext(filename)[0] + '_router.docx'
                output_docx_path = os.path.join(OUTPUT_FOLDER, output_filename)
                
                conversion_success = TestwithSeverity.extract_and_style_audit_findings(pdf_path, output_docx_path)
                
                if conversion_success:
                    message = 'Router/Switch report finished successfully!'
                    download_link = url_for('download_file', filename=output_filename)
                else:
                    message = 'An error occurred during conversion.'
            except Exception as e:
                message = f"An unexpected error occurred: {str(e)}"

    return redirect(url_for('convert_router_page', message=message, download_link=download_link))

@app.route('/convert_firewall', methods=['POST'])
@user_required
def convert_firewall():
    message = None
    download_link = None
    
    if 'pdf_file' not in request.files:
        message = 'No file part in the request.'
    else:
        file = request.files['pdf_file']
        if file.filename == '' or not file.filename.endswith('.pdf'):
            message = 'Please select a PDF file.'
        else:
            try:
                filename = secure_filename(file.filename)
                pdf_path = os.path.join(UPLOAD_FOLDER, filename)
                file.save(pdf_path)
                
                output_filename = os.path.splitext(filename)[0] + '_firewall.docx'
                output_docx_path = os.path.join(OUTPUT_FOLDER, output_filename)
                
                generator = FWADPTE.ReportGenerator(pdf_path, output_docx_path)
                conversion_success = generator.run()
                
                if conversion_success:
                    message = 'Firewall report finished successfully!'
                    download_link = url_for('download_file', filename=output_filename)
                else:
                    message = 'An error occurred during conversion.'
            except Exception as e:
                message = f"An unexpected error occurred: {str(e)}"

    return redirect(url_for('convert_firewall_page', message=message, download_link=download_link))

@app.route('/download/<filename>')
@user_required
def download_file(filename):
    file_path = os.path.join(OUTPUT_FOLDER, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True, download_name=filename)
    else:
        return "File not found.", 404

@app.route('/admin_login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        # Use environment variables with fallback defaults
        admin_username = os.environ.get('ADMIN_USERNAME', 'admin')
        admin_password = os.environ.get('ADMIN_PASSWORD', 'admin_2025')
            
        if username == admin_username and password == admin_password:
            session['admin_logged_in'] = True
            return redirect(url_for('admin'))
        else:
            return render_template('admin_login.html', error='Invalid credentials')
    return render_template('admin_login.html')

@app.route('/admin', methods=['GET', 'POST'])
@admin_required
def admin():
    conn = get_db_connection()
    if not conn:
        return "Database connection failed.", 500
        
    cursor = conn.cursor()

    if request.method == 'POST':
        action = request.form.get('action')
        name = request.form.get('name')
        severity = request.form.get('severity')
        search_type = request.form.get('search_type')
        submitted_by = request.form.get('submitted_by')

        if action == 'approve':
            if search_type == 'firewall':
                cursor.execute('INSERT INTO firewall_vulnerabilities (name, severity) VALUES (%s, %s) ON CONFLICT(name) DO UPDATE SET severity=%s', (name, severity, severity))
            else:
                cursor.execute('INSERT INTO vulnerabilities (name, severity) VALUES (%s, %s) ON CONFLICT(name) DO UPDATE SET severity=%s', (name, severity, severity))
            conn.commit()
        
        cursor.execute('DELETE FROM pending WHERE name = %s', (name,))
        conn.commit()

    cursor.execute("SELECT * FROM pending WHERE search_type = 'router'")
    router_pending = cursor.fetchall()
    
    cursor.execute("SELECT * FROM pending WHERE search_type = 'firewall'")
    firewall_pending = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template('admin.html', router_pending=router_pending, firewall_pending=firewall_pending)

@app.route('/admin/manage_users', methods=['GET', 'POST'])
@admin_required
def manage_users():
    conn = get_db_connection()
    if not conn:
        return "Database connection failed.", 500
    cursor = conn.cursor()
    message = None

    if request.method == 'POST':
        action = request.form.get('action')
        username = request.form.get('username')
        password = request.form.get('password')

        if action == 'add':
            if username and password:
                try:
                    hashed_password = generate_password_hash(password)
                    cursor.execute("INSERT INTO users (username, password_hash, status) VALUES (%s, %s, 'enabled')", (username, hashed_password))
                    conn.commit()
                    message = f"User '{username}' added successfully."
                except pymysql.Error as err:
                    message = f"Error: {str(err)}"
            else:
                message = "Username and password cannot be empty."
        elif action == 'delete':
            # Also remove user from all groups
            cursor.execute("DELETE FROM user_groups WHERE user_id = (SELECT id FROM users WHERE username = %s)", (username,))
            cursor.execute("DELETE FROM users WHERE username = %s", (username,))
            conn.commit()
            message = f"User '{username}' deleted."
        elif action == 'change_password':
            new_password = request.form.get('new_password')
            if username and new_password:
                hashed_password = generate_password_hash(new_password)
                cursor.execute("UPDATE users SET password_hash = %s WHERE username = %s", (hashed_password, username))
                conn.commit()
                message = f"Password for user '{username}' changed successfully."
            else:
                message = "Username and new password cannot be empty."
        elif action == 'toggle_status':
            if username:
                cursor.execute("SELECT status FROM users WHERE username = %s", (username,))
                current_status = cursor.fetchone()
                if current_status:
                    new_status = 'disabled' if current_status['status'] == 'enabled' else 'enabled'
                    cursor.execute("UPDATE users SET status = %s WHERE username = %s", (new_status, username))
                    conn.commit()
                    message = f"User '{username}' status changed to '{new_status}'."
                else:
                    message = f"User '{username}' not found."
            else:
                message = "Username cannot be empty."
        elif action == 'assign_group':
            group_id = request.form.get('group_id')
            if username and group_id:
                try:
                    # Get user ID
                    cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
                    user_result = cursor.fetchone()
                    if user_result:
                        user_id = user_result['id']
                        # Check if user is already in this group
                        cursor.execute("SELECT 1 FROM user_groups WHERE user_id = %s AND group_id = %s", (user_id, group_id))
                        if not cursor.fetchone():
                            cursor.execute("INSERT INTO user_groups (user_id, group_id) VALUES (%s, %s)", (user_id, group_id))
                            conn.commit()
                            # Get group name for message
                            cursor.execute("SELECT name FROM `groups` WHERE id = %s", (group_id,))
                            group_name = cursor.fetchone()['name']
                            message = f"User '{username}' assigned to group '{group_name}' successfully."
                        else:
                            message = f"User '{username}' is already in the selected group."
                    else:
                        message = f"User '{username}' not found."
                except pymysql.Error as err:
                    message = f"Error assigning group: {str(err)}"
            else:
                message = "Username and group selection are required."
        elif action == 'remove_group':
            group_id = request.form.get('group_id')
            if username and group_id:
                try:
                    # Get user ID
                    cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
                    user_result = cursor.fetchone()
                    if user_result:
                        user_id = user_result['id']
                        cursor.execute("DELETE FROM user_groups WHERE user_id = %s AND group_id = %s", (user_id, group_id))
                        conn.commit()
                        # Get group name for message
                        cursor.execute("SELECT name FROM `groups` WHERE id = %s", (group_id,))
                        group_result = cursor.fetchone()
                        group_name = group_result['name'] if group_result else 'Unknown'
                        message = f"User '{username}' removed from group '{group_name}' successfully."
                    else:
                        message = f"User '{username}' not found."
                except pymysql.Error as err:
                    message = f"Error removing from group: {str(err)}"
            else:
                message = "Username and group selection are required."

    # Get all users with their groups (MySQL compatible)
    cursor.execute("SELECT id, username, status FROM users ORDER BY username")
    users_data = cursor.fetchall()
    
    users = []
    for user_data in users_data:
        user = {
            'id': user_data['id'],
            'username': user_data['username'],
            'status': user_data['status'],
            'groups': []
        }
        
        # Get groups for this user
        cursor.execute("""
            SELECT g.id, g.name, g.description 
            FROM `groups` g
            JOIN user_groups ug ON g.id = ug.group_id 
            WHERE ug.user_id = %s
            ORDER BY g.name
        """, (user_data['id'],))
        groups = cursor.fetchall()
        user['groups'] = groups
        users.append(user)
    
    # Get all available groups
    cursor.execute("SELECT id, name, description FROM `groups` ORDER BY name")
    groups = cursor.fetchall()
    
    cursor.close()
    conn.close()
    return render_template('manage_users.html', users=users, groups=groups, message=message)


@app.route('/admin/manage_router_vulnerabilities', methods=['GET', 'POST'])
@admin_required
def manage_router_vulnerabilities():
    conn = get_db_connection()
    if not conn:
        return "Database connection failed.", 500
        
    cursor = conn.cursor()
    
    if request.method == 'POST':
        action = request.form.get('action')
        name = request.form.get('name')

        if action == 'delete':
            cursor.execute('DELETE FROM vulnerabilities WHERE name = %s', (name,))
            conn.commit()
        elif action == 'update':
            new_severity = request.form.get('severity')
            cursor.execute('UPDATE vulnerabilities SET severity = %s WHERE name = %s', (new_severity, name))
            conn.commit()

    cursor.execute('SELECT * FROM vulnerabilities ORDER BY severity DESC')
    vulnerabilities = cursor.fetchall()
    
    cursor.close()
    conn.close()
    
    return render_template('manage_router_vulnerabilities.html', vulnerabilities=vulnerabilities)

@app.route('/admin/manage_firewall_vulnerabilities', methods=['GET', 'POST'])
@admin_required
def manage_firewall_vulnerabilities():
    conn = get_db_connection()
    if not conn:
        return "Database connection failed.", 500
        
    cursor = conn.cursor()
    
    if request.method == 'POST':
        action = request.form.get('action')
        name = request.form.get('name')

        if action == 'delete':
            cursor.execute('DELETE FROM firewall_vulnerabilities WHERE name = %s', (name,))
            conn.commit()
        elif action == 'update':
            new_severity = request.form.get('severity')
            cursor.execute('UPDATE firewall_vulnerabilities SET severity = %s WHERE name = %s', (new_severity, name))
            conn.commit()

    cursor.execute('SELECT * FROM firewall_vulnerabilities ORDER BY severity DESC')
    vulnerabilities = cursor.fetchall()
    
    cursor.close()
    conn.close()
    
    return render_template('manage_firewall_vulnerabilities.html', vulnerabilities=vulnerabilities)

@app.route('/admin/manage_groups', methods=['GET', 'POST'])
@admin_required
def manage_groups():
    conn = get_db_connection()
    if not conn:
        return "Database connection failed.", 500
        
    cursor = conn.cursor()
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'create':
            name = request.form.get('name')
            description = request.form.get('description')
            cursor.execute('INSERT INTO `groups` (name, description) VALUES (%s, %s)', (name, description))
            conn.commit()
        
        elif action == 'delete':
            group_id = request.form.get('group_id')
            cursor.execute('DELETE FROM `groups` WHERE id = %s', (group_id,))
            conn.commit()
            
        elif action == 'update':
            group_id = request.form.get('group_id')
            name = request.form.get('name')
            description = request.form.get('description')
            cursor.execute('UPDATE `groups` SET name = %s, description = %s WHERE id = %s', (name, description, group_id))
            conn.commit()

    # Get all groups with permission counts
    cursor.execute("""
        SELECT g.*, COUNT(gp.permission_id) as permission_count 
        FROM `groups` g 
        LEFT JOIN group_permissions gp ON g.id = gp.group_id 
        GROUP BY g.id, g.name, g.description, g.created_at
        ORDER BY g.name
    """)
    groups = cursor.fetchall()
    
    # Get all permissions for the form
    cursor.execute('SELECT * FROM permissions ORDER BY name')
    permissions = cursor.fetchall()
    
    cursor.close()
    conn.close()
    
    return render_template('manage_groups.html', groups=groups, permissions=permissions)

@app.route('/admin/manage_permissions', methods=['GET', 'POST'])
@admin_required
def manage_permissions():
    conn = get_db_connection()
    if not conn:
        return "Database connection failed.", 500
        
    cursor = conn.cursor()
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'create':
            name = request.form.get('name')
            description = request.form.get('description')
            resource = request.form.get('resource')
            action_name = request.form.get('action_name')
            cursor.execute('INSERT INTO permissions (name, description, resource, action) VALUES (%s, %s, %s, %s)', 
                         (name, description, resource, action_name))
            conn.commit()
        
        elif action == 'delete':
            permission_id = request.form.get('permission_id')
            cursor.execute('DELETE FROM permissions WHERE id = %s', (permission_id,))
            conn.commit()

    cursor.execute('SELECT * FROM permissions ORDER BY resource, name')
    permissions = cursor.fetchall()
    
    cursor.close()
    conn.close()
    
    return render_template('manage_permissions.html', permissions=permissions)

@app.route('/admin/assign_group_permissions', methods=['POST'])
@admin_required
def assign_group_permissions():
    conn = get_db_connection()
    if not conn:
        return jsonify({'ok': False, 'error': 'Database connection failed'}), 500
        
    cursor = conn.cursor()
    
    group_id = request.form.get('group_id')
    permission_ids = request.form.getlist('permission_ids')
    
    # Remove existing permissions for this group
    cursor.execute('DELETE FROM group_permissions WHERE group_id = %s', (group_id,))
    
    # Add new permissions
    for permission_id in permission_ids:
        cursor.execute('INSERT INTO group_permissions (group_id, permission_id) VALUES (%s, %s)', 
                      (group_id, permission_id))
    
    conn.commit()
    cursor.close()
    conn.close()
    
    return redirect(url_for('manage_groups'))

@app.route('/admin/assign_user_groups', methods=['POST'])
@require_permission('manage_users')
def assign_user_groups():
    conn = get_db_connection()
    if not conn:
        return jsonify({'ok': False, 'error': 'Database connection failed'}), 500
        
    cursor = conn.cursor()
    
    user_id = request.form.get('user_id')
    group_ids = request.form.getlist('group_ids')
    
    # Remove existing group assignments for this user
    cursor.execute('DELETE FROM user_groups WHERE user_id = %s', (user_id,))
    
    # Add new group assignments
    for group_id in group_ids:
        cursor.execute('INSERT INTO user_groups (user_id, group_id) VALUES (%s, %s)', 
                      (user_id, group_id))
    
    conn.commit()
    cursor.close()
    conn.close()
    
    return redirect(url_for('manage_users'))

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('user_login'))

def _clean_session_data():
    """Cleans up the temporary data folder and session."""
    session_id = session.pop('manual_review_session_id', None)
    findings_filepath = session.pop('findings_filepath', None)
    if findings_filepath and os.path.exists(findings_filepath):
        os.remove(findings_filepath)
    if session_id:
        session_folder = os.path.join(app.config['TEMP_DATA_FOLDER'], session_id)
        if os.path.exists(session_folder):
            shutil.rmtree(session_folder)

# We are removing type hints here to fix the NameError
def find_section_bbox(page, header_rect, start_keyword, end_keywords):
    """Finds the bounding box of a section on a page for image cropping."""
    try:
        words = page.get_text("words")
        start_word = next((w for w in words if start_keyword.lower() in w[4].lower() and w[1] > header_rect.y1), None)
        if not start_word:
            return None
        end_word = None
        min_end_y = float('inf')
        for keyword in end_keywords:
            candidates = [w for w in words if keyword.lower() in w[4].lower() and w[1] > start_word[1]]
            if candidates:
                first_candidate = min(candidates, key=lambda w: w[1])
                if first_candidate[1] < min_end_y:
                    min_end_y = first_candidate[1]
                    end_word = first_candidate
        page_rect = page.rect
        x0 = page_rect.x0 + 30
        top = start_word[1] - 10
        x1 = page_rect.x1 - 30
        bottom = end_word[1] - 10 if end_word else page_rect.y1 - 30
        if bottom <= top:
            bottom = page_rect.y1 - 30
        if end_word and (end_word[1] - start_word[1] < 20):
            return None
        return fitz.Rect(x0, top, x1, bottom)
    except Exception:
        return None

def process_cis_benchmark(pdf_path):
    """Parses a PDF for findings, adds a 'status' field, and returns them."""
    try:
        doc = fitz.Document(pdf_path)
        full_text = "".join(page.get_text() for page in doc)
    except Exception as e:
        print(f"ERROR: Failed to open or read PDF file: {e}")
        return []

    findings = []
    # Updated regex to better match CIS benchmark structure - looks for numbered sections
    header_regex = re.compile(r"(\d+\.\d+(?:\.\d+)?\s+[^\n]+)", re.MULTILINE)
    matches = list(header_regex.finditer(full_text))

    for i, current_match in enumerate(matches):
        header_text = current_match.group(1).strip().replace('\n', ' ')
        content_end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        # Skip non-finding headers (like table of contents, overview sections)
        if any(skip_word in header_text.lower() for skip_word in ['table of contents', 'terms of use', 'overview', 'contents', 'page ']):
            continue
            
        # Find content for this finding (from end of header to next header)
        content_start_pos = current_match.end()
        content_block = full_text[content_start_pos:content_end_pos]
        # Extract sections with more flexible patterns for CIS benchmarks  
        desc_match = re.search(r"Description:\s*(.*?)(?=\n\s*(?:Rationale|Profile\s+Applicability|Audit|Remediation):|$)", content_block, re.DOTALL | re.IGNORECASE)
        remed_match = re.search(r"Remediation:\s*(.*?)(?=\n\s*(?:Default\s+Value|References|Impact|Additional):|$)", content_block, re.DOTALL | re.IGNORECASE)
        has_audit_section = "Audit:" in content_block
        if not (desc_match or remed_match or has_audit_section):
            continue

        finding = {
            'header': header_text,
            'description': desc_match.group(1).strip() if desc_match else "Not found.",
            'remediation': remed_match.group(1).strip() if remed_match else "Not found.",
            'audit_image_url': None,
            'audit_status': "[Audit section not found or empty.]",
            'status': 'pending'
        }

        if has_audit_section:
            for page in doc:
                search_header = header_text.splitlines()[0].strip()
                header_rects = page.search_for(search_header)
                if header_rects:
                    end_keywords = ["Remediation:", "Default Value:", "References:", "CIS Controls:", "Additional Information:"]
                    bbox = find_section_bbox(page, header_rects[0], "Audit:", end_keywords)
                    if bbox:
                        pix = page.get_pixmap(clip=bbox, dpi=150)
                        image_data = pix.tobytes("png")
                        base64_image = base64.b64encode(image_data).decode('utf-8')
                        finding['audit_image_url'] = f"data:image/png;base64,{base64_image}"
                        finding['audit_status'] = "Image successfully extracted."
                        break
        findings.append(finding)
    doc.close()
    return findings

@app.route('/manual_review', methods=['GET', 'POST'])
@user_required
def manual_review():
    if request.method == 'POST' and 'pdf_file' in request.files:
        file = request.files['pdf_file']
        if file.filename != '' and file.filename.endswith('.pdf'):
            try:
                _clean_session_data()
                session_id = str(uuid.uuid4())
                session['manual_review_session_id'] = session_id
                os.makedirs(os.path.join(app.config['TEMP_DATA_FOLDER'], session_id), exist_ok=True)

                filename = secure_filename(file.filename)
                pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(pdf_path)
                
                all_findings_raw = process_cis_benchmark(pdf_path)
                
                data_filepath = os.path.join(app.config['TEMP_DATA_FOLDER'], session_id, 'findings.json')
                with open(data_filepath, 'w') as f:
                    json.dump(all_findings_raw, f)
                
                session['findings_filepath'] = data_filepath
                session['current_index'] = 0
                
                if not all_findings_raw:
                    return render_template('manual_review.html', message="No valid findings found in the document.", is_empty=True)

                os.remove(pdf_path)
                # Redirect to show findings with navigation (one at a time)
                return redirect(url_for('manual_review'))

            except Exception as e:
                return render_template('manual_review.html', error=f"An error occurred: {str(e)}")
    
    if 'findings_filepath' in session:
        data_filepath = session.get('findings_filepath')
        if not os.path.exists(data_filepath):
             _clean_session_data()
             return render_template('manual_review.html', error="Session data not found. Please re-upload the file.")

        with open(data_filepath, 'r') as f:
            all_findings = json.load(f)

        action = request.args.get('action')
        if action in ['pass', 'fail']:
            try:
                index_to_update = int(request.args.get('index'))
                all_findings[index_to_update]['status'] = action
                with open(data_filepath, 'w') as f:
                    json.dump(all_findings, f)
                
                if index_to_update < len(all_findings) - 1:
                    return redirect(url_for('manual_review', direction='next'))
            except (ValueError, IndexError):
                pass

        current_index = session.get('current_index', 0)
        direction = request.args.get('direction')

        if direction == 'next':
            current_index = min(current_index + 1, len(all_findings) - 1)
        elif direction == 'prev':
            current_index = max(current_index - 1, 0)
        
        session['current_index'] = current_index
        finding = all_findings[current_index]

        return render_template('manual_review.html', 
                               finding=finding, 
                               index=current_index, 
                               total=len(all_findings))
    
    _clean_session_data()
    return render_template('manual_review.html')

def _safe_read_json(filepath, max_retries=3):
    """Safely read JSON file with file locking and retry mechanism."""
    for attempt in range(max_retries):
        try:
            with open(filepath, 'r') as f:
                lock_file(f.fileno(), "r")  # Shared lock for reading
                data = json.load(f)
                lock_file(f.fileno(), unlock_file)  # Unlock
                return data
        except (json.JSONDecodeError, IOError) as e:
            if attempt == max_retries - 1:
                raise e
            # Brief wait before retry
            import time
            time.sleep(0.1 * (attempt + 1))
    return None

def _safe_write_json(filepath, data, max_retries=3):
    """Safely write JSON file using atomic write with file locking."""
    for attempt in range(max_retries):
        try:
            # Create temporary file in same directory for atomic write
            temp_dir = os.path.dirname(filepath)
            with tempfile.NamedTemporaryFile(
                mode='w', 
                dir=temp_dir, 
                delete=False, 
                suffix='.tmp'
            ) as temp_file:
                lock_file(temp_file.fileno(), "w")  # Exclusive lock
                json.dump(data, temp_file, indent=2)
                temp_file.flush()
                os.fsync(temp_file.fileno())  # Ensure write to disk
                temp_filepath = temp_file.name
            
            # Atomic replace
            os.replace(temp_filepath, filepath)
            return True
            
        except (IOError, OSError) as e:
            # Clean up temp file if it exists
            if 'temp_filepath' in locals() and os.path.exists(temp_filepath):
                try:
                    os.unlink(temp_filepath)
                except:
                    pass
            
            if attempt == max_retries - 1:
                raise e
            
            # Brief wait before retry
            import time
            time.sleep(0.1 * (attempt + 1))
    
    return False

@app.route('/manual_review/action', methods=['POST'])
@user_required
def manual_review_action():
    """AJAX endpoint for handling Pass/Fail/Next/Previous actions without page reload"""
    try:
        # Validate session data
        if 'findings_filepath' not in session:
            return jsonify({'ok': False, 'error': 'Session expired. Please re-upload.'}), 400
        
        data_filepath = session.get('findings_filepath')
        if not data_filepath or not os.path.exists(data_filepath):
            _clean_session_data()
            return jsonify({'ok': False, 'error': 'Session data missing. Please re-upload.'}), 400

        # Parse and validate request data
        payload = request.get_json(silent=True)
        if not payload:
            return jsonify({'ok': False, 'error': 'Invalid request: no JSON payload received'}), 400
            
        action = (payload.get('action') or '').lower()
        if not action:
            return jsonify({'ok': False, 'error': 'Invalid request: action is required'}), 400
            
        try:
            idx = int(payload.get('index')) if 'index' in payload else session.get('current_index', 0)
        except (TypeError, ValueError):
            return jsonify({'ok': False, 'error': 'Invalid index: must be a valid number'}), 400

        # Safely read findings file with locking
        try:
            all_findings = _safe_read_json(data_filepath)
            if not all_findings:
                return jsonify({'ok': False, 'error': 'No findings data available'}), 400
        except json.JSONDecodeError as e:
            return jsonify({'ok': False, 'error': f'Data corruption detected. Please re-upload your file. Error: {str(e)}'}), 400
        except Exception as e:
            return jsonify({'ok': False, 'error': f'Error reading findings data: {str(e)}'}), 500

        total = len(all_findings)
        if total == 0:
            return jsonify({'ok': False, 'error': 'No findings available in the dataset'}), 400

        # Validate and adjust index
        idx = max(0, min(idx, total - 1))
        
        # Handle different actions
        if action in ('pass', 'fail'):
            try:
                all_findings[idx]['status'] = action
                success = _safe_write_json(data_filepath, all_findings)
                if not success:
                    return jsonify({'ok': False, 'error': 'Failed to save finding status. Please try again.'}), 500
                
                # Automatically advance to next finding after marking
                idx = min(idx + 1, total - 1)
                
            except Exception as e:
                return jsonify({'ok': False, 'error': f'Error saving finding status: {str(e)}'}), 500
                
        elif action == 'next':
            idx = min(idx + 1, total - 1)
        elif action == 'prev':
            idx = max(idx - 1, 0)
        elif action == 'finished':
            return jsonify({'ok': True, 'action': 'finished', 'redirect': url_for('review_summary')})
        else:
            return jsonify({'ok': False, 'error': f'Invalid action: {action}. Valid actions are: pass, fail, next, prev, finished'}), 400

        # Update session
        session['current_index'] = idx
        
        # Get current finding data
        if idx >= len(all_findings):
            return jsonify({'ok': False, 'error': 'Finding index out of range'}), 400
            
        finding = all_findings[idx]
        
        # Return success response with finding data
        return jsonify({
            'ok': True,
            'action': action,
            'index': idx,
            'total': total,
            'has_prev': idx > 0,
            'has_next': idx < total - 1,
            'is_last': idx == total - 1,
            'finding': {
                'header': finding.get('header', 'No header available'),
                'description': finding.get('description', 'No description available'),
                'remediation': finding.get('remediation', 'No remediation available'),
                'audit_image_url': finding.get('audit_image_url'),
                'audit_status': finding.get('audit_status', 'No audit status available'),
                'status': finding.get('status', 'pending'),
            }
        })
        
    except Exception as e:
        # Comprehensive error handler - always return JSON
        print(f"ERROR in manual_review_action: {str(e)}")
        return jsonify({
            'ok': False, 
            'error': f'An unexpected error occurred: {str(e)}. Please refresh the page and try again.',
            'error_type': type(e).__name__
        }), 500

@app.route('/generate_manual_report')
@user_required
def generate_manual_report():
    if 'findings_filepath' not in session:
        return redirect(url_for('manual_review'))

    data_filepath = session.get('findings_filepath')
    with open(data_filepath, 'r') as f:
        all_findings = json.load(f)

    failed_findings = [f for f in all_findings if f.get('status') == 'fail']

    if not failed_findings:
        _clean_session_data()
        return render_template('convert_router.html', message="No failed findings were recorded. Report not generated.")

    doc = Document()
    doc.add_heading('Manual Review Failed Findings Report', level=1)

    for finding in failed_findings:
        doc.add_heading(finding['header'], level=2)
        doc.add_paragraph(f"Description: {finding['description']}")
        doc.add_paragraph(f"Remediation: {finding['remediation']}")
        doc.add_paragraph()

    output_filename = f"Manual_Review_Report_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
    doc.save(output_path)
    
    _clean_session_data()

    download_link = url_for('download_file', filename=output_filename)
    return render_template('convert_router.html', message='Manual review report finished successfully!', download_link=download_link)



@app.route('/review_summary')
@user_required
def review_summary():
    """Display summary of passed and failed findings"""
    if 'findings_filepath' not in session:
        return redirect(url_for('manual_review'))
        
    data_filepath = session.get('findings_filepath')
    if not os.path.exists(data_filepath):
        _clean_session_data()
        return render_template('manual_review.html', error="Session data not found. Please re-upload the file.")
        
    with open(data_filepath, 'r') as f:
        all_findings = json.load(f)
    
    # Add index to each finding for easy identification
    for idx, finding in enumerate(all_findings):
        finding['index'] = idx
    
    # Separate findings by status
    passed_findings = [f for f in all_findings if f.get('status') == 'pass']
    failed_findings = [f for f in all_findings if f.get('status') == 'fail']
    
    return render_template('review_summary.html', 
                         passed_findings=passed_findings,
                         failed_findings=failed_findings,
                         total_findings=len(all_findings))

@app.route('/review_summary/change_status', methods=['POST'])
@user_required  
def change_status_summary():
    """Change finding status from review summary page"""
    try:
        data = request.get_json()
        finding_index = data.get('index')
        new_status = data.get('status')
        
        if finding_index is None or new_status not in ['pass', 'fail']:
            return jsonify({'ok': False, 'error': 'Invalid parameters'}), 400
        
        if 'findings_filepath' not in session:
            return jsonify({'ok': False, 'error': 'No active session'}), 400
        
        data_filepath = session.get('findings_filepath')
        if not os.path.exists(data_filepath):
            return jsonify({'ok': False, 'error': 'Session data not found'}), 400
        
        # Thread-safe file operations with locking
        lock_file = data_filepath + '.lock'
        max_attempts = 10
        attempt = 0
        
        while attempt < max_attempts:
            try:
                # Try to acquire file lock
                with open(lock_file, 'x') as lock:
                    lock.write(str(os.getpid()))
                break
            except FileExistsError:
                time.sleep(0.1)
                attempt += 1
        
        if attempt >= max_attempts:
            return jsonify({'ok': False, 'error': 'Could not acquire file lock'}), 500
        
        try:
            # Load current findings
            with open(data_filepath, 'r') as f:
                all_findings = json.load(f)
            
            # Validate index
            if finding_index < 0 or finding_index >= len(all_findings):
                return jsonify({'ok': False, 'error': 'Invalid finding index'}), 400
            
            # Update status
            all_findings[finding_index]['status'] = new_status
            
            # Write back atomically  
            temp_filepath = data_filepath + '.tmp'
            with open(temp_filepath, 'w') as f:
                json.dump(all_findings, f, indent=2)
            
            # Atomic move
            os.replace(temp_filepath, data_filepath)
            
            # Calculate new stats
            passed_count = sum(1 for f in all_findings if f.get('status') == 'pass')
            failed_count = sum(1 for f in all_findings if f.get('status') == 'fail')
            
            return jsonify({
                'ok': True, 
                'new_status': new_status,
                'passed_count': passed_count,
                'failed_count': failed_count,
                'finding_header': all_findings[finding_index].get('header', 'Unknown Finding')
            })
            
        finally:
            # Always release lock
            try:
                os.unlink(lock_file)
            except FileNotFoundError:
                pass
                
    except Exception as e:
        return jsonify({'ok': False, 'error': f'Server error: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
