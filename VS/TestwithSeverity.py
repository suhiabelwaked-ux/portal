import re
import pypdf
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import mysql.connector

# Global variable to store submitter name
SUBMITTER_NAME = "suhaib"

def set_submitter_name(name):
    """Set the submitter name for database operations"""
    global SUBMITTER_NAME
    SUBMITTER_NAME = name
    print(f"Submitter name set to: {SUBMITTER_NAME}")

# --- Database Connection and Severity Lookup ---
def get_severity_from_db(finding_title):
    """
    Connects to the MySQL database and fetches the severity for a given finding title.
    It performs a case-insensitive search. If not found, it's added to a pending table.
    """
    try:
        conn = mysql.connector.connect(
            user='root',
            password='admin_2025',
            host='localhost',
            database='vulnerability_db'
        )
        cursor = conn.cursor(dictionary=True)
        # Convert finding_title to lowercase for a case-insensitive match
        cursor.execute("SELECT severity FROM vulnerabilities WHERE LOWER(name) = LOWER(%s)", (finding_title,))
        result = cursor.fetchone()

        # If the vulnerability is not found in the main table
        if not result:
            print(f"Vulnerability '{finding_title}' not found. Adding to pending table.")
            
            # Set default values for severity and search_type
            default_severity = "Unassigned"
            search_type = "router"
            sanitized_title = finding_title.strip()
            
            # Insert the new pending vulnerability
            cursor.execute("INSERT INTO pending (name, severity, search_type, submitted_by) VALUES (%s, %s, %s, %s)", (sanitized_title, default_severity, search_type, SUBMITTER_NAME))
            conn.commit()

            return None

        return result['severity']

    except mysql.connector.Error as err:
        print(f"Database error: {err}")
        return None
    finally:
        if 'conn' in locals() and conn.is_connected():
            cursor.close()
            conn.close()

# --- Utility Functions for Text and Document Styling ---
def clean_text(text):
    """Removes extra whitespace and known text artifacts from a string."""
    if not text:
        return ""
    cleaned_text = ' '.join(text.strip().split())
    cleaned_text = cleaned_text.replace("Type: Best Practice Finding ID: NSA-PRTCL-019", "")
    return cleaned_text.strip()

def style_paragraph(paragraph, font_name='Aptos', size=11, bold=False, spacing_after=0, font_color=RGBColor(0, 0, 0), style=None):
    """Applies styling to a paragraph."""
    if style:
        paragraph.style = style
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = font_color

        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)

    paragraph.paragraph_format.space_after = Pt(spacing_after)

def style_main_heading(doc, text):
    """Adds and styles a main section heading."""
    heading = doc.add_heading(text, level=1)
    run = heading.runs[0]
    font = run.font
    font.size = Pt(14.5)
    font.name = 'Aptos SemiBold'
    font.bold = True
    font.color.rgb = RGBColor(37, 55, 65)
    heading.paragraph_format.space_after = Pt(6)

def write_finding_to_doc(doc, finding, count):
    """Writes a finding into the Word document with formatting."""
    header = finding['header']
    finding_text = clean_text(finding['finding_text'])
    reco_text = clean_text(finding['reco_text'])

    if not (header and finding_text and reco_text):
        if header and header not in globals().get('skipped_sections', []):
            globals().setdefault('skipped_sections', []).append(header)
        return False

    # --- Replacement step ---
    finding_text = finding_text.replace('Nipper', 'FutureTEC')
    reco_text = reco_text.replace('Nipper', 'FutureTEC')

    # Remove unnecessary trailing phrases
    phrases_to_truncate = [
        "These are detailed below.",
        "Those filter rule lists are detailed below"
    ]
    for phrase in phrases_to_truncate:
        if phrase in finding_text:
            finding_text = finding_text.split(phrase)[0].strip()

    title_only = re.sub(r"^\s*\d+(\.\d+)*\s*", "", header).strip()

    if count > 1:
        doc.add_paragraph("")

    title_para = doc.add_paragraph()
    title_para.add_run(f"{count}-  {title_only}")
    style_paragraph(title_para, font_name='Aptos SemiBold', size=12, bold=True, font_color=RGBColor(37, 55, 65), style='Heading 2')

    # --- Severity block ---
    sc_heading = doc.add_paragraph()
    sc_heading.paragraph_format.left_indent = Inches(0.25)
    sc_heading.paragraph_format.space_after = Pt(0)
    font_name = 'Aptos SemiBold'
    font_size = Pt(11)

    run = sc_heading.add_run(f"{count}.1 Severity: ")
    run.font.name = font_name
    run.font.size = font_size
    run.bold = True
    run.font.color.rgb = RGBColor(37, 55, 65)

    # Check if the finding title exists in our database
    specific_severity = get_severity_from_db(title_only)

    severity_colors = {
        'Critical': RGBColor(192, 0, 0),
        'High': RGBColor(255, 0, 0),
        'Medium': RGBColor(255, 192, 0),
        'Low': RGBColor(146, 205, 80)
    }

    if specific_severity and specific_severity in severity_colors:
        # If found, add only the specific severity with its color
        r = sc_heading.add_run(specific_severity)
        r.font.name = font_name
        r.font.size = font_size
        r.bold = True
        r.font.color.rgb = severity_colors[specific_severity]
    else:
        # If not found, show all severities for manual selection
        all_severities = [
            ('Critical', RGBColor(192, 0, 0)),
            ('High', RGBColor(255, 0, 0)),
            ('Medium', RGBColor(255, 192, 0)),
            ('Low', RGBColor(146, 205, 80))
        ]
        for i, (label, color) in enumerate(all_severities):
            if i > 0:
                sep = sc_heading.add_run(' / ')
                sep.font.name = font_name
                sep.font.size = font_size
                sep.bold = True
            r = sc_heading.add_run(label)
            r.font.name = font_name
            r.font.size = font_size
            r.bold = True
            r.font.color.rgb = color


    # --- Description section ---
    desc_heading = doc.add_paragraph(f"{count}.2 Description:")
    desc_heading.paragraph_format.left_indent = Inches(0.25)
    style_paragraph(desc_heading, font_name='Aptos SemiBold', size=11, bold=True, font_color=RGBColor(37, 55, 65))
    desc_heading.paragraph_format.space_after = Pt(0)

    if "FutureTEC determined" in finding_text:
        before, after = finding_text.split("FutureTEC determined", 1)

        first_para = doc.add_paragraph(style='List Bullet')
        first_para.paragraph_format.left_indent = Inches(0.75)
        first_para.add_run(before.strip())
        style_paragraph(first_para, font_name='Aptos', size=11)

        second_para = doc.add_paragraph(style='List Bullet')
        second_para.paragraph_format.left_indent = Inches(0.75)
        second_para.add_run("FutureTEC determined " + after.strip())
        style_paragraph(second_para, font_name='Aptos', size=11)
    else:
        content_para_desc = doc.add_paragraph(style='List Bullet')
        content_para_desc.paragraph_format.left_indent = Inches(0.75)
        content_para_desc.add_run(finding_text)
        style_paragraph(content_para_desc, font_name='Aptos', size=11)

    # --- Recommendation section ---
    reco_heading = doc.add_paragraph(f"{count}.3 Recommendation:")
    reco_heading.paragraph_format.left_indent = Inches(0.25)
    style_paragraph(reco_heading, font_name='Aptos SemiBold', size=11, bold=True, font_color=RGBColor(37, 55, 65))

    if "FutureTEC recommends that:" in reco_text:
        marker_text = "FutureTEC recommends that:"
        before, after = reco_text.split(marker_text, 1)

        # Write the text that came BEFORE the marker text, if it exists
        if before.strip():
            before_para = doc.add_paragraph(style='List Bullet')
            before_para.paragraph_format.left_indent = Inches(0.75)
            before_para.add_run(before.strip())
            style_paragraph(before_para, font_name='Aptos', size=11)

        # Write the marker text itself on a new line
        marker_para = doc.add_paragraph(style='List Bullet')
        marker_para.paragraph_format.left_indent = Inches(0.75)
        marker_para.add_run(marker_text)
        style_paragraph(marker_para, font_name='Aptos', size=11)

        # Split remaining text by semicolon
        after_parts = [p.strip() for p in after.split(";") if p.strip()]
        for part in after_parts:
            bullet_para = doc.add_paragraph()
            bullet_para.paragraph_format.left_indent = Inches(0.75)
            bullet_para.add_run("\u006F  " + part)  # Unicode o
            style_paragraph(bullet_para, font_name='Aptos', size=11)

    else:
        content_para_reco = doc.add_paragraph(style='List Bullet')
        content_para_reco.paragraph_format.left_indent = Inches(0.75)
        content_para_reco.add_run(reco_text)
        style_paragraph(content_para_reco, font_name='Aptos', size=11)

    return True

def extract_and_style_audit_findings(pdf_path, output_docx_path):
    """
    Reads a PDF file from a given path, extracts audit findings,
    and styles them into a Word document.
    Returns True for success, False for failure.
    """
    try:
        print(f"Reading PDF: {pdf_path}...")
        pdf_reader = pypdf.PdfReader(pdf_path)
        full_text = ""
        for page in pdf_reader.pages:
            full_text += page.extract_text() + "\n"

        start_marker = "2 Security Audit"
        header_pattern = re.compile(r"^\s*(2\.\d+)\s+.*$")
        finding_subheader_pattern = re.compile(r"^\s*\d\.\d+\.\d+\s+Finding")
        reco_subheader_pattern = re.compile(r"^\s*\d\.\d+\.\d+\s+Recommendation")
        stop_finding_subheader_pattern = re.compile(r"^\s*\d\.\d+\.\d+\s+(Impact|Ease)")

        sidebar_metadata_pattern = re.compile(r"\s*(Overall:|Impact:|Ease:|Fix:|Type:|Finding ID:).*", re.IGNORECASE)

        junk_line_pattern = re.compile(
            r"Table\s+\d+|detailed in Table|Go to the report contents|Notes\s+for|Severity|Host\s+IP|Port\s*/\s*Protocol|Vulnerability|^\s*Line\s+Access\s+Login|Interface\s+Active\s+Description"
        )

        doc = Document()

        in_audit_section = False
        is_reading_finding = False
        is_reading_reco = False

        current_finding_data = {'header': '', 'finding_text': '', 'reco_text': ''}
        findings = []
        globals().setdefault('skipped_sections', [])

        print("Starting line-by-line processing...")
        lines = full_text.splitlines()

        for line in lines:
            if not in_audit_section:
                if start_marker in line:
                    in_audit_section = True
                continue

            if header_pattern.match(line):
                if current_finding_data.get('header'):
                    findings.append(current_finding_data)
                current_finding_data = {'header': clean_text(line), 'finding_text': '', 'reco_text': ''}
                is_reading_finding, is_reading_reco = False, False
                continue

            if re.search(junk_line_pattern, line):
                is_reading_finding = False
                is_reading_reco = False
                continue

            cleaned_line = sidebar_metadata_pattern.sub("", line).strip()
            if not cleaned_line:
                continue

            if stop_finding_subheader_pattern.match(line):
                is_reading_finding = False
                continue

            if reco_subheader_pattern.match(line):
                is_reading_finding = False
                is_reading_reco = True
                continue

            if finding_subheader_pattern.match(line):
                is_reading_finding = True
                is_reading_reco = False
                continue

            if is_reading_finding:
                current_finding_data['finding_text'] += cleaned_line + " "
            elif is_reading_reco:
                current_finding_data['reco_text'] += cleaned_line + " "

        if current_finding_data.get('header'):
            findings.append(current_finding_data)

        style_main_heading(doc, 'Security Audit')

        written_count = 1
        for finding in findings:
            if write_finding_to_doc(doc, finding, written_count):
                written_count += 1

        doc.save(output_docx_path)
        print(f"\n✅ Successfully created final Word document: {output_docx_path}")
        return True

    except FileNotFoundError:
        print(f"❌ Error: The file '{pdf_path}' was not found.")
        return False
    except Exception as e:
        print(f"❌ An unexpected error occurred: {e}")
        return False

# This __main__ block is for local testing and will not be executed when the module is imported by Flask.
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(
        description="Extracts and styles security audit findings from a PDF into a Word document."
    )
    parser.add_argument("pdf_path", help="The full path to the input PDF file.")
    parser.add_argument(
        "-o",
        "--output",
        help="The path for the output Word document. Defaults to the same name as the input PDF with a .docx extension.",
    )
    parser.add_argument(
        "-u",
        "--user",
        help="The username to associate with pending submissions.",
        default="suhaib"
    )
    args = parser.parse_args()

    if args.output:
        output_path = args.output
    else:
        pdf_directory = os.path.dirname(args.pdf_path)
        pdf_basename = os.path.splitext(os.path.basename(args.pdf_path))[0]
        output_path = os.path.join(pdf_directory, f"{pdf_basename}.docx")

    # Set the submitter name before processing
    set_submitter_name(args.user)
    extract_and_style_audit_findings(args.pdf_path, output_path)